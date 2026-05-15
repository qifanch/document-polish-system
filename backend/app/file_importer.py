from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any


SUPPORTED_EXTENSIONS = {".txt", ".text", ".docx"}


def normalize_text(text: str) -> str:
    lines = [line.rstrip() for line in (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    return "\n".join(lines).strip()


def count_characters(text: str) -> int:
    return len((text or "").replace(" ", "").replace("\n", "").replace("\t", ""))


def parse_imported_file(filename: str, content: bytes) -> dict[str, Any]:
    original_filename = Path(filename or "imported-file").name
    extension = Path(original_filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("仅支持导入 txt、docx 文件，暂不支持 PDF。")
    if not content:
        raise ValueError("导入文件为空，请选择包含正文内容的文件。")

    if extension in {".txt", ".text"}:
        text = parse_txt(content)
        file_type = "txt"
    else:
        text = parse_docx(content)
        file_type = "docx"

    normalized = normalize_text(text)
    if not normalized:
        raise ValueError("未能从文件中提取到文本内容。")

    return {
        "filename": original_filename,
        "fileType": file_type,
        "text": normalized,
        "charCount": count_characters(normalized),
    }


def parse_txt(content: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("文本文件编码无法识别，请转换为 UTF-8 后重试。")


def parse_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("后端缺少 python-docx 依赖，暂时无法解析 docx 文件。") from exc

    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise ValueError("docx 文件解析失败，请确认文件未损坏。") from exc

    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    return "\n".join(parts)
